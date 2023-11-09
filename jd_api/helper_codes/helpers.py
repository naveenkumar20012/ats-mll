"""
Helper Functions for inference.py module.
"""
# Import Necessary Libraries
from __future__ import print_function

import ast
import logging as logger
import os
import re
import subprocess
import tempfile
import unicodedata
from operator import itemgetter
from unicodedata import normalize
from pathlib import Path
import pandas as pd
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
import fitz  # this is pymupdf
import docx
import phonenumbers
import uuid

import sys
from pathlib import Path

sys.path.append(str(Path(__file__).parent))

from utils.logging_helpers import info_logger, error_logger

logger.basicConfig(level="DEBUG")

load_dotenv("./constants.env")

# list of all Hardcoded or fixed variables.
JD_TITE_MAX_IND = int(os.getenv("JD_TITE_MAX_IND"))
WORKEXP_WORDS = ast.literal_eval(os.getenv("WORKEXP_WORDS"))
DESTINATION_FOLDER = os.getenv("DESTINATION_FOLDER")
name_Api = os.getenv("NAME_API")
# Only for work Experience
WORD_INT_DICT = ast.literal_eval(os.getenv("WORD_INT_DICT"))
# keeps only alphanumeric characters in text
KEEP_ALPHANUM = re.compile(r"[\W_]+", re.UNICODE)
# Add spaces before and after special characters
# \] - to match ] , \\ - to match \
# cant keep "." because salary has percentages eg: 0.15% etc
DETECT_PUNCTS = re.compile(r"(\.[,()!/:\\\]])")
DETECT_PUNCTS_SALARY = re.compile(r"([()!/:\\\]])")  # no comma, fullstop

EMPLOYMENTTYPE_FILE = (
    str(Path(__file__).parent.parent) + "/helper_files/employment_type.txt"
)
INDUSTRYTYPE_FILE = (
    str(Path(__file__).parent.parent) + "/helper_files/industry_type.txt"
)
LANGUAGES_FILE = (
    str(Path(__file__).parent.parent) + "/helper_files/languages.txt"
)
LOCATIONS_FOLDER = (
    str(Path(__file__).parent.parent) + "/helper_files/Locations"
)
BENEFITS_PERKS_FILE = (
    str(Path(__file__).parent.parent) + "/helper_files/benefits.txt"
)
EDUCATIONALQUAL_FILE = (
    str(Path(__file__).parent.parent) + "/helper_files/edu_qual.txt"
)

TRUE_LABEL_NAMES = {
    "Job_Title": "job_title",
    "Employment_Type": "employment_type",
    "Tech_Skills": "skills",
    "Edu_Qual": "education",
    "Work_Exp": "work_experience",
    "Location": "location",
    "Benefits_perks": "benefits_perks",
    "Annual_Salary_Range": "salary_range",
    "Monthly_Salary_Range": "salary_range",
    "Industry_Type": "industry_type",
}

# Exhaustive list for Employment Type
with open(EMPLOYMENTTYPE_FILE, "r") as emp_str:
    EMPLOYMENTTYPE_LST = ast.literal_eval(emp_str.read())

# Exhaustive list for Industry Type
with open(INDUSTRYTYPE_FILE, "r") as ind_str:
    INDUSTRYTYPE_LST = ast.literal_eval(ind_str.read())

# Exhaustive list for Languages
with open(LANGUAGES_FILE, "r", encoding="utf-8") as f:
    LANG_LST = f.read()
LANG_LST = ast.literal_eval(LANG_LST)
LANG_LST = [
    unicodedata.normalize("NFD", lang.lower())
    .encode("ascii", "ignore")
    .decode("utf-8")
    for lang in LANG_LST
]

# Exhaustive list for Locations
ALL_LOC_FILES = os.listdir(LOCATIONS_FOLDER)
ALL_LOCS = []
for filename in ALL_LOC_FILES:
    if filename.endswith(".csv"):
        file_path = os.path.join(LOCATIONS_FOLDER, filename)
        # info_logger(file_path)
        df = pd.read_csv(file_path)
        locs = df["Locations"].tolist()
        # info_logger(len(locs))
        locs_new = []
        for loc in locs:
            loc = loc.lower().strip()
            loc = KEEP_ALPHANUM.sub("", loc)
            loc = DETECT_PUNCTS.sub("", loc)
            locs_new.append(loc)
        ALL_LOCS.extend(locs_new)

# Exhaustive list for Benefits and Perks
with open(BENEFITS_PERKS_FILE, "r") as ben_str:
    BENEFITS_LST = ast.literal_eval(ben_str.read())

# Exhaustive list for Educational Qualification
with open(EDUCATIONALQUAL_FILE, "r") as edu_str:
    EDU_QUAL_LST = ast.literal_eval(edu_str.read())


def get_true_label_name(label):
    if TRUE_LABEL_NAMES.get(label):
        return TRUE_LABEL_NAMES.get(label)
    else:
        return label


# Functions required for data loading (JD's in text file or pdf or doc or docx)

# https://towardsdatascience.com/extracting-headers-and-paragraphs-from-pdf-using-pymupdf-\
# 676e8421c467


def fonts(doc, granularity=False):
    """Extracts fonts and their usage in PDF documents.
    :param doc: PDF document to iterate through
    :type doc: <class 'fitz.fitz.Document'>
    :param granularity: also use 'font', 'flags' and 'color' to discriminate text
    :type granularity: bool
    :rtype: [(font_size, count), (font_size, count}], dict
    :return: most used fonts sorted by count, font style information
    """
    styles = {}
    font_counts = {}

    for page in doc:
        blocks = page.getText("dict")["blocks"]
        for block in blocks:  # iterate through the text blocks
            if block["type"] == 0:  # block contains text
                for line in block["lines"]:  # iterate through the text lines
                    for span in line[
                        "spans"
                    ]:  # iterate through the text spans
                        if granularity:
                            identifier = "{0}_{1}_{2}_{3}".format(
                                span["size"],
                                span["flags"],
                                span["font"],
                                span["color"],
                            )
                            styles[identifier] = {
                                "size": span["size"],
                                "flags": span["flags"],
                                "font": span["font"],
                                "color": span["color"],
                            }
                        else:
                            identifier = "{0}".format(span["size"])
                            styles[identifier] = {
                                "size": span["size"],
                                "font": span["font"],
                            }
                        # count the fonts usage
                        font_counts[identifier] = (
                            font_counts.get(identifier, 0) + 1
                        )

    font_counts = sorted(font_counts.items(), key=itemgetter(1), reverse=True)

    if len(font_counts) < 1:
        raise ValueError("Zero discriminating fonts found!")

    return font_counts, styles


def font_tags(font_counts, styles):
    """Returns dictionary with font sizes as keys and tags as value.
    :param font_counts: (font_size, count) for all fonts occurring in document
    :type font_counts: list
    :param styles: all styles found in the document
    :type styles: dict
    :rtype: dict
    :return: all element tags based on font-sizes
    """
    p_style = styles[
        font_counts[0][0]
    ]  # get style for most used font by count (paragraph)
    p_size = p_style["size"]  # get the paragraph's size

    # sorting the font sizes high to low, so that we can append the right integer to each tag
    font_sizes = []
    for (font_size, _) in font_counts:
        font_sizes.append(float(font_size))
    font_sizes.sort(reverse=True)

    # aggregating the tags for each font size
    idx = 0
    size_tag = {}
    for size in font_sizes:
        idx += 1
        if size == p_size:
            idx = 0
            size_tag[size] = "<p>"
        if size > p_size:
            size_tag[size] = "<h{0}>".format(idx)
        elif size < p_size:
            size_tag[size] = "<s{0}>".format(idx)

    return size_tag


def headers_para(doc, size_tag):
    """
    Scrapes headers & paragraphs from PDF and return texts with element tags.
    :param doc: PDF document to iterate through
    :type doc: <class 'fitz.fitz.Document'>
    :param size_tag: textual element tags for each size
    :type size_tag: dict
    :rtype: list
    :return: texts with pre-prended element tags
    """
    header_para = []  # list with headers and paragraphs
    first = True  # boolean operator for first header
    previous_s = {}  # previous span

    for page in doc:
        blocks = page.getText("dict")["blocks"]
        for block in blocks:  # iterate through the text blocks
            if block["type"] == 0:  # this block contains text

                # REMEMBER: multiple fonts and sizes are possible IN one block

                block_string = ""  # text found in block
                for line in block["lines"]:  # iterate through the text lines
                    for span in line[
                        "spans"
                    ]:  # iterate through the text spans
                        if span["text"].strip():  # removing whitespaces:
                            # info_logger(s)
                            if first:
                                previous_s = span
                                first = False
                                block_string = (
                                    size_tag[span["size"]] + span["text"]
                                )
                            else:
                                if span["size"] == previous_s["size"]:

                                    if block_string and all(
                                        (c == "|") for c in block_string
                                    ):
                                        # block_string only contains pipes
                                        block_string = (
                                            size_tag[span["size"]]
                                            + span["text"]
                                        )
                                    if block_string == "":
                                        # new block has started, so append size tag
                                        block_string = (
                                            size_tag[span["size"]]
                                            + span["text"]
                                        )
                                    else:  # in the same block, so concatenate strings
                                        block_string += " " + span["text"]

                                else:
                                    header_para.append(block_string)
                                    block_string = (
                                        size_tag[span["size"]] + span["text"]
                                    )

                                previous_s = span

                    # new block started, indicating with a pipe
                    block_string += "|"

                header_para.append(block_string)

    return header_para


def pdf_to_html(doc):
    text = ""
    for i, page in enumerate(doc):
        text += page.getText("html")
    return text


############ DOCX EXTRACT JD HTML FROM  #############
def get_para_run_html(para):
    html = ""
    for run in para.runs:
        if run.text:
            bold = run.bold
            italic = run.italic
            underline = run.underline
            if bold:
                html += "<b>"
            if italic:
                html += "<i>"
            if underline:
                html += "<u>"
            html += run.text
            if underline:
                html += "</u>"
            if italic:
                html += "</i>"
            if bold:
                html += "</b>"
    return html


def get_docx_html(doc):
    whole_html = ""
    list_initialised = False
    for para in doc.paragraphs:
        html = ""
        if para.text:
            if "List" in para.style.name:
                if not list_initialised:
                    html += "<ul>"
                    list_initialised = True
                html += "<p>"
                html += "<li>"
                html += get_para_run_html(para)
                html += "</li>"
            else:
                if list_initialised:
                    html += "</ul>"
                    list_initialised = False
                html += "<p>"
                html += get_para_run_html(para)

            html += "</p>"
        whole_html += html
    return whole_html


############################################################


def clean_jd_list(jd_lst):
    """
    Remove all unicode character, strip extra spaces in font and n in the end for each sentence.

    Args:
        jd_lst(list): list of all sentences in JD.
    Returns:
        Cleaned list by removing unnecessary characters.
    """
    text_lst_clean = []
    for text in jd_lst:
        # strip extra spaces in fornt and \n in the end
        text = (text.strip()).rstrip("\n").replace("\u2019", "'")
        # check if we have , or . between numbers
        if bool(
            re.search(r"[0-9]+,[0-9]+", text)
            or re.search(r"[0-9]+\.[0-9]+", text)
        ):
            info_logger("salary", text)
            text = DETECT_PUNCTS_SALARY.sub(" \\1 ", text)
        else:
            text = DETECT_PUNCTS.sub(
                " \\1 ", text
            )  # add space before and after special characters.
        text = text.replace("\t", "")  # remove tab spaces
        text = text.replace("  ", " ")  # replace double spaces with one space
        text = normalize("NFKD", text)  # Remove all unicode character
        # Remove all unicode character and covert to utf-8
        text = text.encode("ascii", "ignore").decode("utf-8")
        if not text == "":
            text_lst_clean.append(text)
        # text_lst_clean.append(text)
    return text_lst_clean


def jd_txt(filepath):
    """
    Load text file as list with each datapoint as a line.

    Args:
        filepath(str): Absolute path of the JD text file

    Returns: list of all sentences in a JD text file
    """
    with open(filepath, "r", encoding="utf-8") as txt_f:
        text_lst = txt_f.readlines()

    text_lst = clean_jd_list(text_lst)
    return text_lst


def jd_pdf(filepath):
    """
    Load pdf file as list with each datapoint as a line.

    Args:
        filepath(str): Absolute path of the JD file

    Returns: list of all sentences in a JD file
    """
    pdf = fitz.open(filepath)
    font_counts, styles = fonts(pdf)
    size_tag = font_tags(font_counts, styles)
    header_para = headers_para(pdf, size_tag)

    ############ PDF HTML VERSION AND CLEANING ############
    html_version = pdf_to_html(pdf)
    img_cleaner = re.compile("<img[^<]+?>")
    cleantext = re.sub(img_cleaner, "", html_version)
    para_cleaner = re.compile('<p style=".*?">')
    cleantext = re.sub(para_cleaner, '<p style="">', cleantext)
    hexcode_cleaner = re.compile("&#x.*?;")
    cleantext = re.sub(hexcode_cleaner, " ", cleantext)

    html = BeautifulSoup(cleantext, "lxml")
    for ele in html.find_all("p"):
        if len(ele.get_text(strip=True)) == 0:
            # Remove empty tag
            ele.extract()

    # Removing the width and height characters in the JD
    if html:
        html = str(html)
        re_width = r"width:[0-9]+pt;"
        re_height = r"height:[0-9]+pt;"
        html = re.sub(re_width, "", html)
        html = re.sub(re_height, "", html)
    ##########################################################

    pdf_text = []
    for block in header_para:
        block_new = (
            block.replace("<h1>", "")
            .replace("<h2>", "")
            .replace("<p>", "")
            .replace("|", "")
            .replace("<h5>", "")
            .replace("<s1>", "")
        )
        pdf_text.append(block_new)
    pdf_text = clean_jd_list(pdf_text)
    return pdf_text, str(html).replace("\n", "")


def jd_docx(filepath):
    """
    Load file as list with each datapoint as a line.

    Args:
        filepath(str): Absolute path of the JD file

    Returns: list of all sentences in a JD file
    """
    doc = docx.Document(filepath)
    full_text = []
    html_version = get_docx_html(doc)

    for para in doc.paragraphs:
        full_text.append(para.text)

    fulltext_lst = clean_jd_list(full_text)

    # Removing the width and height characters in the JD
    if html_version:
        html_version = str(html_version)
        re_width = r"width:[0-9]+pt;"
        re_height = r"height:[0-9]+pt;"
        html_version = re.sub(re_width, "", html_version)
        html_version = re.sub(re_height, "", html_version)

    return fulltext_lst, html_version


def doc_to_docx_conversion(filepath):
    """
    Convert doc file to docx file.

    Args:
        filepath(str): Absolute path of the JD file

    Returns: path of converted docx file
    """
    new_filename = ".".join(filepath.split("/")[-1].split(".")[:-1]) + ".docx"
    new_filepath = os.path.join(DESTINATION_FOLDER, new_filename)
    os.environ["PATH"] += ":/bin"
    test = subprocess.Popen(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "docx",
            filepath,
            "--outdir",
            DESTINATION_FOLDER,
        ],
        stdout=subprocess.PIPE,
    )
    _ = test.communicate()[0]
    return new_filepath


def jd_doc(filepath):
    """
    Load doc file, convert to docx and then load docx file
    as list with each datapoint as a line.

    Args:
        filepath(str): Absolute path of the JD file

    Returns: list of all sentences in a JD file
    """
    new_path = doc_to_docx_conversion(filepath)
    doc_text_lst, html_version = jd_docx(new_path)
    return doc_text_lst, html_version


# Helper Functions related to model
def iob_to_biluo(tags):
    out = []
    tags = list(tags)
    while tags:
        out.extend(_consume_os(tags))
        out.extend(_consume_ent(tags))
    return out


def _consume_os(tags):
    while tags and tags[0] == "O":
        yield tags.pop(0)


def _consume_ent(tags):
    if not tags:
        return []
    tag = tags.pop(0)
    target_in = "I" + tag[1:]
    target_last = "L" + tag[1:]
    length = 1
    while tags and tags[0] in {target_in, target_last}:
        length += 1
        tags.pop(0)
    label = tag[2:]
    if length == 1:
        if len(label) == 0:
            raise ValueError(f"Ill-formed IOB input detected: {tag}")
        return ["U-" + label]
    else:
        start = "B-" + label
        end = "L-" + label
        middle = [f"I-{label}" for _ in range(1, length - 1)]
        return [start] + middle + [end]


def tags_to_entities(tags):
    """Note that the end index returned by this function is inclusive.
    To use it for Span creation, increment the end by 1."""
    entities = []
    start = None
    for i, tag in enumerate(tags):
        if tag is None or tag.startswith("-"):
            # TODO: We shouldn't be getting these malformed inputs. Fix this.
            if start is not None:
                start = None
            else:
                entities.append(("", i, i))
        elif tag.startswith("O"):
            pass
        elif tag.startswith("I"):
            if start is None:
                raise ValueError(
                    f"Invalid BILUO tag sequence: Got a tag starting with 'I' \n without a preceding 'B' (beginning of an entity). \n Tag sequence:\n{tags[: i + 1]}"
                )
        elif tag.startswith("U"):
            entities.append((tag[2:], i, i))
        elif tag.startswith("B"):
            start = i
        elif tag.startswith("L"):
            if start is None:
                raise ValueError(
                    f"Invalid BILUO tag sequence: Got a tag starting with 'L' \n without a preceding 'B' (beginning of an entity). \n Tag sequence:\n{tags[: i + 1]}"
                )
            entities.append((tag[2:], start, i))
            start = None
        else:
            raise ValueError(f"Invalid BILUO tag: '{tag}'.")
    return entities


def predicted_json_manipulation(ind, pred_json, dic_final):
    labs = list(pred_json.keys())

    pred_json_new = {}

    for lab in labs:  # for each label type
        content = pred_json[
            lab
        ]  # get predicted words info of that label in datapoint
        pred_json_new = post_process_functions(
            ind, pred_json_new, content, lab
        )

    #     #remove empty dictionaries because no prediction happened there.
    #     if pred_json_new == {}:
    #         continue

    for key, vals in pred_json_new.items():
        work_ex_label = get_true_label_name("Work_Exp")
        annual_salary_label = get_true_label_name("Annual_Salary_Range")
        monthly_salary_label = get_true_label_name("Monthly_Salary_Range")
        for val in vals:
            if key in dic_final.keys():
                # info_logger("val:", val)
                if key == work_ex_label:
                    dic_final[key].append([val[2], val[-1]])
                elif key == annual_salary_label or key == monthly_salary_label:
                    continue
                else:
                    dic_final[key].append(val[2])
            else:
                if key == work_ex_label:
                    dic_final[key] = [[val[2], val[-1]]]
                elif key == annual_salary_label or key == monthly_salary_label:
                    dic_final[key] = val[4]
                else:
                    dic_final[key] = [val[2]]

    # To get all the datapoints which were eliminated in post-processing.
    pred_json = dict(sorted(pred_json.items()))
    pred_json_new = dict(sorted(pred_json_new.items()))

    return dic_final


# Helper Functions related to Work Experience
def unique_lower_words(text_lst, words_lst):
    """
    for work_exp. Clean the sentence and detected workexp words.

    Args:
        text_lst(list): list containing sentences.
        words_lst(list): list containing predicted words from sentences.

    Returns:
        unique_words : lowercase and Unique words for each datapoint (sentence).
        text_lst : list of sentences corresponding tho the words predicted.
    """
    if isinstance(text_lst, list) and isinstance(words_lst, list):
        unique_words = []
        text_lst_new = []
        for text, word in zip(text_lst, words_lst):
            word = word.lower()
            if word in unique_words:
                pass
            else:
                unique_words.append(word)
                text_lst_new.append(text)
        return text_lst_new, unique_words
    elif isinstance(text_lst, str) and isinstance(words_lst, str):
        words_lst = words_lst.lower()
        text_lst = text_lst.lower()
        return text_lst, words_lst
    else:
        info_logger(
            "Datatypes of text_lst and words_lst are not list or strings"
        )


def convert_words_to_int(text):
    """
    Used in work experience post process to convert "two" to 2.

    Args:
        text(str): datapoint or sentence which has a word related to workexp.

    Returns:
        res: converted datapoint or sentence which has a word related to workexp.
    """
    res = ""
    for ele in re.split(r" |-", text):
        if ele in WORD_INT_DICT.keys():
            res = res + WORD_INT_DICT[ele] + " "
        else:
            res = res + ele + " "
    return res


# Post-processing functions for each label

#################################### UTIL FUNCTIONS FOR SALARY DETECTION ###############################################

currencies = ["€", "$", "₹"]
indian_monetary_terms = ["lpa", "lakh", "l.p.a.", "inr"]
common_salary_pat = r"\d+(?:\.?,?\d+)*"
common_salary_pat_with_L = r"\d+(?:\.?,?\d+)*L"
percentage_detection_re = r"\d+(?:\.?,?\d+)*%"
salary_pat_with_k = r"\d+(?:\.?,?\d+)*k"

duration_terms_annual = ["year", "/yr", "lpa", "annum", "p.a."]
duration_terms_monthly = ["month", "monthly"]
duration_terms_hourly = ["hourly"]


def get_salary_range_from_str(salary_str):
    """
    Returns the salary in float values in a list
    args: salary_str (str): string to get the salary range from
    return: (list)(float): a list with the max and min of the salary
    """
    salary_range_raw = re.findall(common_salary_pat, salary_str)
    salary_range = list(map(lambda x: x.replace(",", ""), salary_range_raw))
    # convert to float
    salary_range = list(map(lambda x: round(float(x), 2), salary_range))
    return salary_range


def get_duration(salary_str):
    """
    Gets the duration of the salary being mentioned
    args: salary_str (str): string to get the salary range from
    return: (str): either 'monthly', 'hourly', 'annual'
    """
    if any(term in salary_str for term in duration_terms_monthly):
        duration = "monthly"
    elif any(term in salary_str for term in duration_terms_annual):
        duration = "annual"
    elif any(term in salary_str for term in duration_terms_hourly):
        duration = "hourly"
    else:
        duration = "annual"

    return duration


def convert_to_annual(salary_range, duration):
    """
    Converts the salary range to annual salary
    args: salary_str (str): string to get the salary range from
          duration (str): either 'monthly', 'hourly', 'annual'
    return: (list)(float): a list with the max and min of the salary in annual range
    """
    if duration == "monthly":
        # convert to per annum
        salary_range = list(map(lambda x: x * 12, salary_range))
    elif duration == "hourly":
        # Skeptical about this (maybe need to include an additional param of duration in the API itself)
        salary_range = list(map(lambda x: x * 8 * 12, salary_range))
    return salary_range


def detect_currency(salary_str):
    """
    Get the currency from the salary
    args: salary_str (str): string to get the salary range from
    return: (str): the symbol of currency
    """
    for currency in currencies:
        if currency in salary_str:
            return currency

    for term in indian_monetary_terms:
        if term in salary_str:
            return "₹"


def convert_to_specific_range(min, max, duration):
    monthly_range = [num for num in range(5000, 100001, 5000)]
    annual_range = [num for num in range(100000, 10000001, 100000)]
    min_value, max_value = 0, 0
    if duration == "monthly":
        for num_index in range(0, len(monthly_range) - 1):
            if min in range(
                monthly_range[num_index], monthly_range[num_index + 1] + 1
            ):
                min_value = monthly_range[num_index]
            if max in range(
                monthly_range[num_index], monthly_range[num_index + 1] + 1
            ):
                if max == monthly_range[num_index + 1]:
                    max_value = monthly_range[num_index + 1]
                else:
                    max_value = monthly_range[num_index]
    elif duration == "annual":
        for num_index in range(0, len(annual_range) - 1):
            if min in range(
                annual_range[num_index], annual_range[num_index + 1] + 1
            ):
                min_value = annual_range[num_index]
            if max in range(
                annual_range[num_index], annual_range[num_index + 1] + 1
            ):
                max_value = annual_range[num_index]

    return min_value, max_value


#################################################################################################################


def salary_post_process(content):
    """
    Extracts the annual salary from the predicted entity
    args: salary_str (str): string to get the salary range from
    return: (json): min, max salary and the currency
    """
    content_mod = []
    for datapoint in content:
        datapoint_mod = datapoint.copy()
        salary_str = datapoint[2]
        lakh_term = re.findall(common_salary_pat_with_L, salary_str)
        salary_str = salary_str.lower()
        detected_currency = None
        # Detect currency
        detected_currency = detect_currency(salary_str)
        if not detected_currency and lakh_term:
            detected_currency = "₹"

        info_logger("detected_currency", detected_currency)
        info_logger(
            "salary pattern", re.findall(common_salary_pat, salary_str)
        )
        # Remove if there are terms containing %
        percentage_terms = re.findall(percentage_detection_re, salary_str)
        for term in percentage_terms:
            salary_str = salary_str.replace(term, "")

        # Converting to proper salary range
        duration = get_duration(salary_str)
        salary_range = get_salary_range_from_str(salary_str)

        # Changing the salary range values to absolute values
        if detected_currency == "$":
            if re.search(salary_pat_with_k, salary_str):
                salary_range = list(map(lambda x: x * 1000, salary_range))

        elif detected_currency == "₹":
            if "lakh" in salary_str or "lpa" in salary_str or lakh_term:
                salary_range = list(
                    map(lambda x: round(x * 100000, 2), salary_range)
                )
                info_logger("salary", salary_range)
        # info_logger(salary_range)
        if len(salary_range) == 2:
            min_salary, max_salary = sorted(salary_range)
            info_logger("salary min max", min_salary, max_salary)
            min_salary, max_salary = convert_to_specific_range(
                min_salary, max_salary, duration
            )
            salary_info = {
                "min": min_salary,
                "max": max_salary,
                "duration": duration,
            }
            datapoint_mod.append(salary_info)
            content_mod.append(datapoint_mod)
        elif len(salary_range) == 1:
            min_salary, max_salary = salary_range[0], salary_range[0]
            min_salary, max_salary = convert_to_specific_range(
                min_salary, max_salary, duration
            )
            salary_info = {
                "min": min_salary,
                "max": max_salary,
                "duration": duration,
            }
            datapoint_mod.append(salary_info)
            content_mod.append(datapoint_mod)
        else:
            info_logger("Salary wrong prediction: ", datapoint)
            pass

    # info_logger("new_salary:", content_mod)
    return content_mod


def location_post_process(content):
    """
    Post-processing for Location
        1.One datapoint should not have more than 3 locations.(because more than
            3 locations can be recognised as default locations specified)
        2. The detected locations should be present in the exhaustive list of
            locations (not using this currently because we only have
            locations of india)
    Args:
        content(list): predictions from NER model.

    Returns: removed elements from content which doesn't satisfy
            post-processing rules
    """
    # to calculate how many location predicts were present for this label in each
    # sentence(datapoint).
    content_rem = content.copy()
    sent_dict = {}
    for datapoint in content:  # for each prediction
        start, end, word, sent = datapoint
        # making a new dict whose key is the datapoint text and values are info regarding
        # the predicted word for this label
        sent_dict.setdefault(sent, []).append((start, end, word))

    for key, value in sent_dict.items():
        # It contains multiple useless location which are like placeholders.
        if len(value) > 3:
            # if there are more than 3 location predicted in a sentence, we can treat
            # it as unimportant data and remove it from predictions.
            for val in value:
                # If one datapoint contains more than 3 location tags, We will delete those
                # tags from output json content
                content_rem.remove([val[0], val[1], val[2], key])
        else:
            # The detected locations should be in the exhaustive list. This can be a flexible rule
            # (might or might not be used)
            # content_rem = [i if i[2].lower() in ALL_LOCS else info_logger(i[2].lower(), content_rem) \
            # for i in content_rem]
            # info_logger(content_rem)
            # content_rem = [i for i in content_rem if(KEEP_ALPHANUM.sub('', i[2].\
            # lower().strip()) not in ALL_LOCS)]
            # pass
            for val in value:
                if all(
                    loc not in KEEP_ALPHANUM.sub("", val[2].lower().strip())
                    for loc in ALL_LOCS
                ):
                    info_logger("Location not in Locations list: ", val[2])
                    content_rem.remove([val[0], val[1], val[2], key])
                else:
                    pass
    return content_rem


def employmenttype_post_process(content):
    """
    Post-processing for Employment Type
        1. The detected employment should be present in the exhaustive list of
            employmenttype
    Args:
        content(list): predictions from NER model.

    Returns: removed elements from content which doesn't satisfy
            post-processing rules

    """
    # Post-processing for Employment type
    content_rem = content.copy()
    for pred in content:  # for each predicted word for Employment_Type label
        word = pred[2]
        # keeps only alphanumeric characters in text (part of text cleaning process)
        word_alpha = KEEP_ALPHANUM.sub("", word)
        # remove "-" from word (part of text cleaning the word to compare with exhaustive
        # list we have)
        word_alpha = re.sub("-", "", word_alpha)
        # not strict match
        one_token_match_employtype = 0
        for token in word_alpha.lower().split(" "):
            # if the predicted wor is not in exhaustive list of employment_type, remove
            # it from predictions.
            if any(emptype in token.lower() for emptype in EMPLOYMENTTYPE_LST):
                one_token_match_employtype = 1
            else:
                pass
        if one_token_match_employtype == 0:
            content_rem.remove(pred)
    #         #strict match
    #         #if the predicted word is not in exhaustive list of employment_type, remove it
    #         # from predictions.
    #         if word_alpha.lower() not in EMPLOYMENTTYPE_LST:
    #             content_rem.remove(i)
    return content_rem


def industrytype_post_process(content):
    """
    Post-processing for Industry Type
        1. The detected Industry type should be present in the exhaustive list of
            industrytype
    Args:
        content(list): predictions from NER model.

    Returns: removed elements from content which doesn't satisfy
            post-processing rules

    """
    content_rem = content.copy()
    for pred in content:
        word = pred[2]
        # keeps only alphanumeric characters in text (part of text cleaning process)
        word_alpha = KEEP_ALPHANUM.sub("", word)
        # remove "-" from word (part of text cleaning the word to compare with exhaustive
        # list we have)
        word_alpha = re.sub("-", "", word_alpha)
        # not strict match
        one_token_match_indtype = 0
        for token in word_alpha.lower().split(" "):
            # if the predicted wor is not in exhaustive list of employment_type, remove
            # it from predictions.
            if any(indtype in token.lower() for indtype in INDUSTRYTYPE_LST):
                one_token_match_indtype = 1
            else:
                pass
        if one_token_match_indtype == 0:
            content_rem.remove(pred)
    #         #strict match
    #         #if the predicted wor is not in exhaustive list of industry_type, remove it
    #         #from predictions.
    #         if word_alpha.lower() not in INDUSTRYTYPE_LST:
    #             content.remove(i)
    return content_rem


def eduqual_post_process(content):
    """
    Post-processing for Educational Qualification
        1. The detected Educational Qualification should be present in the exhaustive list of
            Educational Qualification
    Args:
        content(list): predictions from NER model.

    Returns: removed elements from content which doesn't satisfy
            post-processing rules

    """
    content_rem = content.copy()
    for pred in content:  # for each predicted word for Employment_Type label
        word = pred[2]
        # keeps only alphanumeric characters in text (part of text cleaning process)
        word_alpha = KEEP_ALPHANUM.sub("", word)
        # remove "-" from word (part of text cleaning the word to compare with exhaustive
        # list we have)
        word_alpha = re.sub("-", "", word_alpha)
        one_token_match_eduqual = 0
        for token in word_alpha.lower().split(" "):
            # if the predicted word is not in exhaustive list of employment_type, remove
            # it from predictions.
            if any(eduqual in token.lower() for eduqual in EDU_QUAL_LST):
                one_token_match_eduqual = 1
            else:
                pass
        if one_token_match_eduqual == 0:
            content_rem.remove(pred)
        else:
            pass
    return content_rem


def benefits_perks_post_process(content):
    """
    Post-processing for Benefits
        1. The detected Benefits should be present in the exhaustive list of
            Benefits
    Args:
        content(list): predictions from NER model.

    Returns: removed elements from content which doesn't satisfy
            post-processing rules

    """
    content_rem = content.copy()
    for pred in content:
        word = pred[2]
        # keeps only alphanumeric characters in text (part of text cleaning process)
        word_alpha = KEEP_ALPHANUM.sub("", word)
        # remove "-" from word (part of text cleaning the word to compare with exhaustive
        # list we have)
        word_alpha = re.sub("-", "", word_alpha)
        # not a strict match
        one_token_match_bens = 0
        for token in word_alpha.lower().split(" "):
            # if the predicted word is not in exhaustive list of employment_type, remove
            # it from predictions.
            if any(ben in token.lower() for ben in BENEFITS_LST):
                one_token_match_bens = 1
            else:
                pass
        if one_token_match_bens == 0:
            content_rem.remove(pred)
        else:
            pass
        # #strict match
        # #if the predicted wor is not in exhaustive list of Benefits_perks, remove
        # # it from predictions.
        # if word_alpha.lower() not in BENEFITS_LST:
        #     info_logger(word_alpha)
        #     content.remove(i)
    return content_rem


def workexp_post_process(content):
    """
    Post-processing for Work Experience
        1.Convert words to ints eg: "two" to "2".
        2.Convert float work experience to int.
        3.Segregate if we have work experience in predicted word,
            range or max experience
        4.Check if sentence has years or yrs or yr in it.
    Args:
        content(list): predictions from NER model.

    Returns: removed elements from content which doesn't satisfy
            post-processing rules

    """
    content_mod = (
        []
    )  # the final modified content after processing words related to work_exp
    for pred in content:  # for each predicted word
        new_pred = pred.copy()
        # convert word and text(datapoint) to lowercase.
        text_lower, word_lower = unique_lower_words(pred[-1], pred[-2])
        # If the word contains any number in text form, convert it into a numeric.Eg: "one" to 1
        word_lower = convert_words_to_int(word_lower)
        non_work_exp_words = ["old", "age"]
        if any(
            non_workexp in text_lower for non_workexp in non_work_exp_words
        ):
            info_logger("Wrong Work Experience detected. : ", pred)
        else:
            exps = re.findall(
                r"\d+(?:\.\d+)?", word_lower
            )  # find all float numbers in the text
            exps = [
                int(float(float_num)) for float_num in exps
            ]  # convert float numbers to intergers.
            if exps == []:  # If no integers are detected in workexp
                info_logger(
                    "No Integers in Work Experience detected. : ", pred
                )
                # content.remove(i)
            # If we have only one year mentioned, we can take it as min experience.
            elif (len(exps) == 1) and (
                any(i in text_lower for i in WORKEXP_WORDS)
            ):
                # if ["+", minimum", "min", "atleast"] in exps:
                min_exp = exps[0]
                max_exp = None

                # append min and max work exp also.
                new_pred.extend([{"min": min_exp, "max": max_exp}])
                # append this new list [st, end, word, text, minexp, maxexp] to new content list
                content_mod.append(new_pred)
            # if we have range or two experiences mentioned, we can make min as lowest value
            # and ,max as the highest value.
            elif (len(exps) == 2) and (
                any(i in text_lower for i in WORKEXP_WORDS)
            ):
                min_exp = sorted(exps)[0]
                max_exp = sorted(exps)[1]
                new_pred.extend([{"min": min_exp, "max": max_exp}])
                content_mod.append(new_pred)
            # If there are more than 3 work experiences in a given datapoint, delete it. (This might
            #  have edge cases. eg: multiple eperiences in single datapoint. But didn't encounter
            # this yet.)
            elif len(exps) > 2:
                info_logger(pred)
                info_logger(
                    "Number of integers or floats in text is greater than 2 : ",
                    len(exps),
                )
                # content.remove(pred)
            else:
                info_logger(pred)
                info_logger(
                    "Datapoint is not satisfying work experience conditions: ",
                    pred,
                )

    return content_mod


def language_post_process(content):
    """
    Post-processing for Language
        1. The detected Language should be present in the exhaustive list of
            Language
    Args:
        content(list): predictions from NER model.

    Returns: removed elements from content which doesn't satisfy
            post-processing rules

    """
    content_rem = content.copy()
    for pred in content:
        # This logic takes cases where our predicted word is "English language".
        remove = 0
        for word in pred[2].lower().split(" "):
            if any(
                lang in word for lang in LANG_LST
            ):  # If word in LANG_LST exhaustive list, keep it. else remove it.
                remove = 1
        if remove == 0:
            content_rem.remove(pred)
        else:
            # in contain at least one word which is language. So, we won't remove it.
            pass
    return content_rem


def jobtitle_post_process(ind, content):
    """
    Post-processing for Job-title
        1. The detected Job-title should be present within first 4 lines of JD
    Args:
        ind(int): Index in which Job title is detected
        content(list): predictions from NER model.

    Returns: removed elements from content which doesn't satisfy
            post-processing rules

    """
    content_rem = content.copy()

    for pred in content:
        if ind < JD_TITE_MAX_IND:
            pass
        else:
            content_rem.remove(pred)

    return content_rem


def post_process_functions(ind, pred_json_new, content, lab):
    """
    Do post-processing for one datapoint.

    Args:
        ind(int) : Index of the datapoint in JD.(used for Job title)
        pred_json_new(dict): predicted dictionary with structure {"label":[st, end, word, sentence]}
        content(list): predictions that were obtained from model [st, end, label, word]
        lab(str) : the name of the label

    Returns:
        processed predicted dictionary with structure {"label":[st, end, word, sentence]}.
            for workexp, structure is {"label":[st, end, word, sentence, min_value, max value]}
    """
    true_label = get_true_label_name(lab)
    info_logger("=>content prediction", lab, content)
    # Post processing for Location
    if lab == "Location":
        content = location_post_process(content)
        # Assigning the modified contents for the label.
        pred_json_new[true_label] = content

    # Post-processing for Employment type
    elif lab == "Employment_Type":
        content = employmenttype_post_process(content)
        pred_json_new[true_label] = content

    # Post-processing for Industry type
    elif (
        lab == "Industry_Type"
    ):  # for each predicted word for Industry_Type label
        content = industrytype_post_process(content)
        # Assigning the modified contents for the label.
        pred_json_new[true_label] = content

    # Post-processing for Educational Qualification
    elif lab == "Edu_Qual":
        content = eduqual_post_process(content)
        pred_json_new[true_label] = content

    # Post-processing for Benefits_Perks
    elif (
        lab == "Benefits_perks"
    ):  # for each predicted word for Benefits_perks label
        content = benefits_perks_post_process(content)
        # Assigning the modified contents for the label.
        pred_json_new[true_label] = content

    # Post-processing for Work Experience
    elif lab == "Work_Exp":
        content = workexp_post_process(content)
        pred_json_new[true_label] = content

    # Post-processing for language
    elif lab == "Language":
        content = language_post_process(content)
        pred_json_new[true_label] = content

    # Post-processing for Job Title
    elif lab == "Job_Title":
        content = jobtitle_post_process(ind, content)
        if content == []:
            pass
        else:
            pred_json_new[true_label] = content

    # Post-processing for Salary
    elif lab == "Annual_Salary_Range" or lab == "Monthly_Salary_Range":
        # info_logger("Salary:", content)
        content = salary_post_process(content)
        pred_json_new[true_label] = content

    else:
        pred_json_new[true_label] = content

    return pred_json_new


# ==================================================================================================================================================================
# QUICK IMPLEMENTATION OF BASIC DETAIL EXTRACTOR FROM RESUMES
# ==================================================================================================================================================================


def extract_name(texts):
    """
    For calling name api and cleaning
    Args:
        texts: whole resume  text

    Returns:
        name
    """
    try:
        names = []
        # in first 10 lines
        name = requests.post(
            name_Api, json={"text": " ".join(texts[:10])}
        ).json()["name"]
        if name != "":
            names.append(name)
        # if not in first 10 line go for other lines
        if len(names) < 1:
            name = requests.post(
                name_Api, json={"text": " ".join(texts[10:])}
            ).json()["name"]
            if name != "":
                names.append(name)

        # name cleaning
        if len(names) > 0:
            names = [re.sub(r"[^A-Za-z .]+", "", name) for name in names]
            return max(names, key=len).title()
        else:
            return None
    except Exception as e:
        error_logger("Name API is not working", e)
        return None


def extract_email(text):
    """
    Helper function to extract email id from text

    :param text: plain text extracted from resume file
    """
    email_regex = r"[\w\.-]+@[\w\.-]+(?:\.[\w]+)+"
    email = re.findall(email_regex, text)
    if email:
        try:
            email_id = email[0].split()[0].strip(";")
            return email_id
        except IndexError:
            return None


def find_phone(text):
    try:
        return list(iter(phonenumbers.PhoneNumberMatcher(text, None)))[
            0
        ].raw_string
    except:
        try:
            return re.search(
                r"(\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4})",
                text,
            ).group()
        except:
            return ""


def get_special_parsed_text(text):
    text = text.split("\n")
    return text


def extract_linkedin(text):
    link = re.findall(r"linkedin.com/[a-z]+/[\S]+", text)
    if link:
        return link[0]


def extract_linkedin_spe(text):
    common_words = [
        ".com",
        ".in",
        " ",
        "profile",
        "education",
        "skills",
        "experience",
        "summary",
        ":",
    ]
    text = get_special_parsed_text(text)
    for index, t in enumerate(text[:-1]):
        link = re.findall(r"linkedin.com/[a-z]+/[\S]+", t)
        if link:
            link = link[0]
            next_line = text[index + 1]
            if all(
                [
                    a.lower() not in next_line.lower().strip()
                    for a in common_words
                ]
            ):
                adjusted_link_1 = link + next_line
                adjusted_link = re.findall(
                    r"linkedin.com/[a-z]+/[\S]+", adjusted_link_1
                )
                return "https://" + adjusted_link[0]
            else:
                return "https://" + link


def extract_github(text):
    link = re.findall(r"github.com/[\S]+", text)
    if link:
        return "https://" + link[0]


def extract_basic_details(text, tagger):
    data = {}
    data["name"] = extract_name(text)
    data["email"] = extract_email(text)
    data["phone_number"] = find_phone(text)
    data["linkedin_url"] = extract_linkedin_spe(text)
    data["github_url"] = extract_github(text)
    return data


def get_text_from_pdf(filepath):
    text = ""
    with fitz.open(filepath) as doc:
        for page in doc:
            text += page.getText()
    return text


def get_text_from_docx(filepath):
    """
    Load file as list with each datapoint as a line.

    Args:
        filepath(str): Absolute path of the JD file

    Returns: list of all sentences in a JD file
    """
    doc = docx.Document(filepath)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    return "\n".join(full_text)


from dataclasses import dataclass


@dataclass
class CoreUtil:
    """Core Util"""

    def get_short_uuid(self):
        """Generate Short UUID."""
        return str(uuid.uuid4()).replace("-", "")[:10].upper()


class TempFileUtil:
    """A Temporary File Utility."""

    def __init__(self, file_format):
        # Create Temp folder
        self.temp_dir = tempfile.mkdtemp()
        # Create unpredictable filename
        self.filename = f"{CoreUtil().get_short_uuid()[:8]}.{file_format}"
        # Create temp file path
        self.filepath = os.path.join(self.temp_dir, self.filename)

    def delete(self):
        # Delete temp file
        if os.path.exists(self.filepath):
            os.remove(self.filepath)
        # Delete temp folder
        if os.path.exists(self.temp_dir):
            os.rmdir(self.temp_dir)

    def __del__(self):
        # Call on destruction
        self.delete()


from google.cloud import storage


def get_gcs_client():
    """Return Google Storage Client."""
    return storage.Client.from_service_account_json(
        "/home/praharsha/gcp_service.json"
    )


def download_blob_from_gcs(blob_name, destination_file):
    """Download the file to specified destination."""
    bucket_name = "ats-backend-new"
    client = get_gcs_client()

    if client:
        bucket = client.get_bucket(bucket_name)
        blob = bucket.blob(blob_name)
        if blob.exists():
            blob.download_to_filename(destination_file)
            return destination_file


def doc_to_pdf_conversion(filepath):
    """
    Convert doc file to docx file.

    Args:
        filepath(str): Absolute path of the JD file

    Returns: path of converted docx file
    """
    new_filename = ".".join(filepath.split("/")[-1].split(".")[:-1]) + ".pdf"
    new_filepath = os.path.join(DESTINATION_FOLDER, new_filename)
    os.environ["PATH"] += ":/bin"
    test = subprocess.Popen(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            filepath,
            "--outdir",
            DESTINATION_FOLDER,
        ],
        stdout=subprocess.PIPE,
    )
    _ = test.communicate()[0]
    return new_filepath


def saving_file_local_ats(filepath):
    """
    Take filepath, i fits a gcs path, return filepath stored
    in local directory.

    Args:
        filepath(str): Absolute path of the file.

    Returns:
        Absolute path of file wrt to local directory.
    """
    filename = filepath.split("/")[-1]
    if filepath.startswith("gs://"):  # that means environment is "prod"

        storage_client = get_gcs_client()

        bucket = storage_client.get_bucket("ats-backend-new")

        # Construct a client side representation of a blob.
        # Note `Bucket.blob` differs from `Bucket.get_blob` as it doesn't retrieve
        # any content from Google Cloud Storage. As we don't need additional data,
        # using `Bucket.blob` is preferred here.
        blob_name = "/".join(filepath.split("/")[3:])
        logger.info(blob_name)
        blob = bucket.blob(blob_name)
        dest_file_path = os.path.join(DESTINATION_FOLDER + filename)
        logger.info(dest_file_path)
        try:
            blob.download_to_filename(dest_file_path)
        except:
            return False

    else:  # If file is in local, no need to copy to local.
        dest_file_path = filepath

    if dest_file_path.endswith(".docx"):
        try:
            dest_file_path = doc_to_pdf_conversion(dest_file_path)
            # dest_file_path = dest_file_path.replace('.docx', '.pdf')
        except Exception as e:
            logger.error(e)

    return dest_file_path


def extract_details_from_resume(resume_filepath, tagger):
    filename_gcs = resume_filepath
    resume_filepath = saving_file_local_ats(filename_gcs)
    if not resume_filepath:
        return {"error": "File not found"}

    # Check file extension format type
    if resume_filepath.endswith(".pdf"):
        # Create Temp file manager object
        temp_file_manager = TempFileUtil("pdf")
    elif resume_filepath.endswith(".docx"):
        temp_file_manager = TempFileUtil("docx")
    elif resume_filepath.endswith(".doc"):
        temp_file_manager = TempFileUtil("doc")

    # Download blob from storage
    # resume_filepath = download_blob_from_gcs(
    #     resume_filepath, temp_file_manager.filepath
    # )
    logger.info(resume_filepath)
    if resume_filepath.lower().endswith(".pdf"):
        resume_text = get_text_from_pdf(resume_filepath)
    elif resume_filepath.lower().endswith(".docx"):
        resume_text = get_text_from_docx(resume_filepath)
    elif resume_filepath.lower().endswith(".doc"):
        resume_filepath = doc_to_docx_conversion(resume_filepath)
        resume_text = get_text_from_docx(resume_filepath)
    else:
        return {"error": "Unsupported file type"}

    try:
        with open(
            f'resumeparser_output/{resume_filepath.split("/")[-1]}.txt', "w"
        ) as resume_output:
            resume_output.write(resume_text)
    except Exception as e:
        logger.info(e)

    candidate_details = extract_basic_details(resume_text, tagger)
    return candidate_details
